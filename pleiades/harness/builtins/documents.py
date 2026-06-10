"""Document tools — read PDFs, Word docs, spreadsheets, and images.

The agent shouldn't be blind to half the file types users actually work with.
All deps are optional; each tool reports exactly what to install if missing.

    read_pdf      pypdf (or pdfplumber if installed — better tables)
    read_docx     python-docx — paragraphs + tables as markdown
    read_xlsx     openpyxl — sheets as markdown tables
    read_image    Pillow metadata; description via Claude vision when a key is set
"""

from __future__ import annotations

import base64
from pathlib import Path

from ..tools import tool

_MAX_CHARS = 40_000


def _check(path: str) -> tuple[Path | None, str]:
    p = Path(path).expanduser()
    if not p.exists():
        return None, f"Error: no such file: {path}"
    if p.is_dir():
        return None, f"Error: {path} is a directory."
    return p, ""


def _clip(text: str, what: str) -> str:
    if len(text) > _MAX_CHARS:
        return text[:_MAX_CHARS] + f"\n... [truncated — {what} continues]"
    return text


# ---------------------------------------------------------------------------
@tool(safe=True, tags=("doc", "read"))
def read_pdf(path: str, start_page: int = 1, end_page: int = 0) -> str:
    """Extract text from a PDF, optionally restricted to a page range.

    path: the .pdf file to read.
    start_page: first page, 1-indexed.
    end_page: last page inclusive (0 = to the end).
    """
    p, err = _check(path)
    if err:
        return err

    # pdfplumber first (better layout/tables), pypdf as the lighter fallback
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(p) as pdf:
            total = len(pdf.pages)
            s, e = _page_range(start_page, end_page, total)
            parts = []
            for i in range(s, e):
                parts.append(f"--- page {i + 1} ---\n"
                             + (pdf.pages[i].extract_text() or "(no text)"))
            return _clip(f"[{total} pages]\n" + "\n".join(parts), "PDF")
    except ImportError:
        pass

    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        return ("Error: PDF support needs a library — "
                "pip install pypdf  (or pdfplumber for better table extraction)")
    reader = PdfReader(str(p))
    total = len(reader.pages)
    s, e = _page_range(start_page, end_page, total)
    parts = [f"--- page {i + 1} ---\n"
             + (reader.pages[i].extract_text() or "(no text)")
             for i in range(s, e)]
    return _clip(f"[{total} pages]\n" + "\n".join(parts), "PDF")


def _page_range(start: int, end: int, total: int) -> tuple[int, int]:
    s = max(0, start - 1)
    e = min(total, end) if end > 0 else total
    return s, max(s, e)


# ---------------------------------------------------------------------------
@tool(safe=True, tags=("doc", "read"))
def read_docx(path: str) -> str:
    """Read a Word document — paragraphs, headings, and tables as markdown.

    path: the .docx file to read.
    """
    p, err = _check(path)
    if err:
        return err
    try:
        import docx  # type: ignore
    except ImportError:
        return "Error: Word support needs a library — pip install python-docx"

    d = docx.Document(str(p))
    parts: list[str] = []
    for para in d.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            parts.append("#" * min(int(level), 6) + " " + text)
        else:
            parts.append(text)
    for i, table in enumerate(d.tables):
        rows = [[c.text.strip().replace("\n", " ") for c in r.cells]
                for r in table.rows]
        if not rows:
            continue
        md = ["| " + " | ".join(rows[0]) + " |",
              "|" + "---|" * len(rows[0])]
        md += ["| " + " | ".join(r) + " |" for r in rows[1:]]
        parts.append(f"[table {i + 1}]\n" + "\n".join(md))
    return _clip("\n\n".join(parts) or "(empty document)", "document")


# ---------------------------------------------------------------------------
@tool(safe=True, tags=("doc", "read"))
def read_xlsx(path: str, sheet: str = "", max_rows: int = 100) -> str:
    """Read a spreadsheet — each sheet rendered as a markdown table.

    path: the .xlsx file to read.
    sheet: a specific sheet name (empty = all sheets).
    max_rows: cap on rows returned per sheet.
    """
    p, err = _check(path)
    if err:
        return err
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return "Error: spreadsheet support needs a library — pip install openpyxl"

    wb = openpyxl.load_workbook(str(p), data_only=True, read_only=True)
    names = [sheet] if sheet else wb.sheetnames
    parts: list[str] = []
    for name in names:
        if name not in wb.sheetnames:
            return f"Error: no sheet '{name}'. Sheets: {', '.join(wb.sheetnames)}"
        ws = wb[name]
        rows = []
        for j, row in enumerate(ws.iter_rows(values_only=True)):
            if j >= max_rows:
                rows.append((f"... [truncated at {max_rows} rows "
                             f"of {ws.max_row}]",))
                break
            rows.append(tuple("" if v is None else str(v) for v in row))
        if not rows:
            parts.append(f"## {name}\n(empty sheet)")
            continue
        width = max(len(r) for r in rows)
        norm = [list(r) + [""] * (width - len(r)) for r in rows]
        md = ["| " + " | ".join(norm[0]) + " |", "|" + "---|" * width]
        md += ["| " + " | ".join(r) + " |" for r in norm[1:]]
        parts.append(f"## {name}\n" + "\n".join(md))
    wb.close()
    return _clip("\n\n".join(parts), "spreadsheet")


# ---------------------------------------------------------------------------
_IMG_MEDIA = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
              ".gif": "image/gif", ".webp": "image/webp"}


@tool(safe=True, tags=("doc", "read"))
def read_image(path: str, question: str = "") -> str:
    """Describe an image. Uses Claude vision when an API key is configured; always returns basic metadata (format, dimensions).

    path: the image file (.png/.jpg/.gif/.webp).
    question: optional — something specific to look for in the image.
    """
    p, err = _check(path)
    if err:
        return err

    meta = f"[{p.name}"
    try:
        from PIL import Image  # type: ignore
        with Image.open(p) as im:
            meta += f" — {im.format}, {im.width}x{im.height}, {im.mode}"
    except ImportError:
        meta += " — pip install Pillow for dimensions"
    except Exception as e:
        return f"Error: cannot open image: {e}"
    meta += f", {p.stat().st_size} bytes]"

    media = _IMG_MEDIA.get(p.suffix.lower())
    if media is None:
        return meta + "\nUnsupported format for vision (use png/jpg/gif/webp)."

    try:
        from ..config import Config, HAIKU
        cfg = Config.load()
        if not cfg.anthropic_api_key:
            import os
            if not os.environ.get("ANTHROPIC_API_KEY"):
                return meta + "\n(no ANTHROPIC_API_KEY — vision description unavailable)"
        import anthropic
        client = anthropic.Anthropic(api_key=cfg.anthropic_api_key or None)
        prompt = question or ("Describe this image precisely: layout, text, "
                              "data, and anything notable. Be concise.")
        resp = client.messages.create(
            model=HAIKU, max_tokens=1024,
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {
                    "type": "base64", "media_type": media,
                    "data": base64.standard_b64encode(p.read_bytes()).decode(),
                }},
                {"type": "text", "text": prompt},
            ]}],
        )
        desc = "".join(b.text for b in resp.content if b.type == "text")
        return meta + "\n" + desc
    except ImportError:
        return meta + "\n(pip install anthropic for vision descriptions)"
    except Exception as e:
        return meta + f"\n(vision call failed: {e})"
